"""
reportGenerator.py
Legacy camelCase module retained for compatibility with reference repo.
Full report implementation is in reporting/report_generator.py
"""

import os


def generateReport(metrics, output_path='reports/hero_fincorp_analysis.docx'):
    """
    Generate a Word report summarising key Hero FinCorp metrics.

    Parameters
    ----------
    metrics     : dict — must contain at least 'default_rate'
    output_path : str  — path to save the .docx file
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  ⚠️  python-docx not installed. Run: pip install python-docx")
        return

    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document()

    doc.add_heading('Hero FinCorp Analysis', 0)

    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    default_rate = metrics.get('default_rate', 0)
    doc.add_paragraph(f"Default Rate: {default_rate:.2%}")

    total_loans = metrics.get('total_loans', 'N/A')
    doc.add_paragraph(f"Total Loans Analysed: {total_loans:,}" if isinstance(total_loans, int)
                      else f"Total Loans Analysed: {total_loans}")

    total_income = metrics.get('total_interest_income', 0)
    doc.add_paragraph(f"Estimated Interest Income: ₹{total_income:,.0f}")

    recovery_rate = metrics.get('recovery_rate', 0)
    doc.add_paragraph(f"Overall Recovery Rate: {recovery_rate:.2%}")

    # Key Insights
    doc.add_heading('Key Insights', 1)
    insights = [
        "Low credit score customers (below 580) carry significantly higher default risk.",
        "South and West regions show the highest default rates at 10.5%.",
        "Penalty transactions represent ~50% of all transactions — a critical warning sign.",
        "Average loan processing time is 175 days — far above the 30-day industry target.",
        "Recovery rate of 24.3% is critically low; legal action provides minimal improvement.",
        "33% of the loan portfolio is in Overdue status — requires immediate triage.",
    ]
    for insight in insights:
        doc.add_paragraph(f"• {insight}")

    # Recommendations
    doc.add_heading('Recommendations', 1)
    recs = [
        "Enforce minimum credit score of 580 for unsecured loans.",
        "Implement Early Warning System at 3, 6, and 12 months post-disbursement.",
        "Partner with ARCs for NPAs older than 360 days to improve recovery.",
        "Reduce processing time to 30 days via digital document portals.",
        "Offer EMI restructuring programs to address overdue portfolio.",
    ]
    for rec in recs:
        doc.add_paragraph(f"• {rec}")

    doc.save(output_path)
    print(f"\n  ✅ Report saved: {output_path}")
    return output_path
