"""
Insight Reporter
Generates a formatted Word (.docx) report consolidating findings
from all 20 Hero FinCorp analysis tasks.
"""
import os
from config import REPORT_PATH, FIGURE_PATH

os.makedirs(REPORT_PATH, exist_ok=True)
os.makedirs(FIGURE_PATH, exist_ok=True)


def generate_full_report(df, task_results=None):
    """
    Build and save the Hero FinCorp analysis report.

    Parameters
    ----------
    df           : master merged DataFrame
    task_results : dict of analysis outputs from each task (optional)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  python-docx not installed — run: pip install python-docx")
        return

    print("\n" + "=" * 55)
    print("  Generating Insight Report...")
    print("=" * 55)

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Colour palette
    PRIMARY   = RGBColor(0xE3, 0x1E, 0x24)
    DARK      = RGBColor(0x1A, 0x1A, 0x2E)
    ACCENT    = RGBColor(0x2E, 0x86, 0xAB)

    # --------------------------------------------------
    # HELPER FUNCTIONS
    # --------------------------------------------------
    def add_heading(text, level=1, colour=None):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if colour:
            for run in h.runs:
                run.font.color.rgb = colour
        return h

    def add_body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        return p

    def add_kv(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        lbl = p.add_run(f"  {label}: ")
        lbl.bold = True
        lbl.font.color.rgb = DARK
        p.add_run(str(value))

    def insert_figure(filename, width=Inches(5.5)):
        path = os.path.join(FIGURE_PATH, filename)
        if os.path.exists(path):
            doc.add_picture(path, width=width)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            add_body(f"[Figure not available: {filename}]")

    def insert_table(dataframe, max_rows=15):
        if dataframe is None or dataframe.empty:
            return
        visible = dataframe.head(max_rows).reset_index()
        cols = visible.columns.tolist()
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = 'Table Grid'
        hdr_cells = tbl.rows[0].cells
        for i, col in enumerate(cols):
            hdr_cells[i].text = str(col)
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E31E24')
            shd.set(qn('w:val'), 'clear')
            tc_pr.append(shd)
        for ri, row in visible.iterrows():
            row_cells = tbl.add_row().cells
            for ci, val in enumerate(row):
                row_cells[ci].text = str(round(val, 4) if isinstance(val, float) else val)
                if ri % 2 == 0:
                    tc_pr = row_cells[ci]._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    shd.set(qn('w:val'), 'clear')
                    tc_pr.append(shd)
        doc.add_paragraph()

    # --------------------------------------------------
    # COVER PAGE
    # --------------------------------------------------
    title_para = doc.add_heading('Hero FinCorp', 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.color.rgb = PRIMARY
        run.font.size = Pt(32)

    subtitle = doc.add_paragraph('Credit Portfolio — Data Analysis Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = DARK

    doc.add_paragraph()
    meta = doc.add_paragraph('20 Analysis Tasks  |  April 2026')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)
    doc.add_page_break()

    # --------------------------------------------------
    # EXECUTIVE SUMMARY
    # --------------------------------------------------
    add_heading('Executive Summary', 1, PRIMARY)
    add_body(
        'This report presents findings from a structured data analysis of Hero FinCorp\'s '
        'credit portfolio covering approximately 70,000 customers, 90,000 loans, '
        '82,600 applications, 9,000 default records, 50 branches, and 495,000 transactions.'
    )

    total_rows   = len(df)
    dflt_rate    = df['DEFAULT_FLAG'].mean() * 100 if 'DEFAULT_FLAG' in df.columns else 0
    disb_total   = df['LOAN_AMOUNT'].sum() / 1e9    if 'LOAN_AMOUNT'  in df.columns else 0
    overdue_pct  = (df['LOAN_STATUS'] == 'Overdue').mean() * 100 if 'LOAN_STATUS' in df.columns else 0

    add_kv('Total Loans Analysed',  f'{total_rows:,}')
    add_kv('Total Disbursed',       f'₹{disb_total:.2f} Billion')
    add_kv('Overall Default Rate',  f'{dflt_rate:.1f}%')
    add_kv('Overdue Portfolio',     f'{overdue_pct:.1f}% of active loans')
    add_kv('Avg Credit Score',      f'{df["CREDIT_SCORE"].mean():.0f}' if 'CREDIT_SCORE' in df.columns else 'N/A')
    add_kv('Avg Interest Rate',     f'{df["INTEREST_RATE"].mean():.2f}%' if 'INTEREST_RATE' in df.columns else 'N/A')
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 1 — Data Quality
    # --------------------------------------------------
    add_heading('Task 1: Data Quality & Preparation', 1, PRIMARY)
    add_body(
        'All six datasets were validated and standardised. Column names normalised to uppercase. '
        'Duplicate records removed; numeric nulls imputed with median, categorical nulls with mode. '
        'Date columns parsed and invalid entries dropped for critical date fields. '
        'Outliers capped via IQR method on LOAN_AMOUNT, INTEREST_RATE, and DEFAULT_AMOUNT. '
        'Domain constraints applied: LOAN_AMOUNT > 0, INTEREST_RATE ≥ 0.'
    )
    add_heading('Dataset Overview', 2)
    quality_rows = [
        ['Dataset',      'Records', 'Columns', 'Key Observation'],
        ['customers',    '70,000',  '14',      'No critical nulls; income and score outliers capped'],
        ['loans',        '90,000',  '12',      '~30k null COLLATERAL_DETAILS — treated as unsecured'],
        ['applications', '82,600',  '10',      '~12k null LOAN_ID for rejected apps; expected'],
        ['defaults',     '9,000',   '9',       '~3k null RECOVERY_STATUS — excluded from rate calc'],
        ['branches',     '50',      '9',       'No issues identified'],
        ['transactions', '495,000', '9',       'No issues identified'],
    ]
    tbl = doc.add_table(rows=len(quality_rows), cols=4)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(quality_rows):
        for ci, val in enumerate(row):
            tbl.rows[ri].cells[ci].text = val
            if ri == 0:
                tbl.rows[ri].cells[ci].paragraphs[0].runs[0].bold = True
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 2
    # --------------------------------------------------
    add_heading('Task 2: Exploratory Data Analysis', 1, PRIMARY)
    add_body(
        'Loan amounts and EMI values show right-skewed distributions with notable clustering. '
        'Credit scores are broadly uniform. Monthly applications and approvals remain stable '
        'over the analysis period. Default rates vary meaningfully by region.'
    )
    for fig in ['task_2_loan_distribution.png', 'task_2_emi_distribution.png',
                'task_2_credit_score.png', 'task_2_region_default.png',
                'task_2_monthly_applications.png', 'task_2_monthly_approved.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 3
    # --------------------------------------------------
    add_heading('Task 3: Default Risk Correlation', 1, PRIMARY)
    add_body(
        'Credit score carries a negative correlation with default flag, confirming its predictive '
        'utility. Interest rate shows a weak positive association with default. '
        'EMI amount, overdue amount, and default amount move together — '
        'indicating that financial stress manifests across multiple indicators simultaneously.'
    )
    for fig in ['task_3_correlation_main.png', 'task_3_correlation_pairwise.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 4
    # --------------------------------------------------
    add_heading('Task 4: Branch & Regional Performance', 1, PRIMARY)
    add_body(
        'Branch delinquency rates vary widely. Disbursement volumes are concentrated '
        'in a subset of high-performing branches. Branch-level linkage to loan data '
        'is unavailable (no BRANCH_ID in loans/applications), limiting granular drill-down.'
    )
    for fig in ['task_4_branch_loan.png', 'task_4_branch_delinquency.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 5
    # --------------------------------------------------
    add_heading('Task 5: Borrower Segmentation', 1, PRIMARY)
    add_body(
        'Customers segmented into Low / Medium / High income tiers and credit score tiers. '
        'High-risk borrowers (low income + low credit score) represent a disproportionate '
        'share of defaults. High-value borrowers (high income + high credit score) '
        'are candidates for premium product cross-sell.'
    )
    for fig in ['task_5_income_segment.png', 'task_5_credit_segment.png', 'task_5_loan_status.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 6
    # --------------------------------------------------
    add_heading('Task 6: Advanced Statistical Analysis', 1, PRIMARY)
    add_body(
        'Extended correlation matrices confirm that lower credit scores and higher interest '
        'rates co-occur with elevated default risk. Pairwise analysis of EMI, recovery rate, '
        'and default amount surfaces the financial co-stress relationship among distressed borrowers.'
    )
    for fig in ['task_6_adv_default_corr.png', 'task_6_adv_pairwise_corr.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 7
    # --------------------------------------------------
    add_heading('Task 7: Payment & Recovery Analysis', 1, PRIMARY)
    add_body(
        'Transaction patterns reveal widespread repayment stress — penalty transactions '
        'constitute a significant share of all activity. Recovery rates vary across regions '
        'and default reasons. Collections efficiency is an area requiring strategic improvement.'
    )
    for fig in ['task_7_txn_type.png', 'task_7_recovery_reason.png', 'task_7_recovery_region.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 8
    # --------------------------------------------------
    add_heading('Task 8: EMI Risk Analysis', 1, PRIMARY)
    add_body(
        'Higher EMI bands carry elevated default probability. Mid-to-high EMI brackets '
        'show the steepest risk gradient. Borrowers with EMI-to-income ratios above 40% '
        'should trigger enhanced monitoring at origination.'
    )
    for fig in ['task_8_emi_default.png', 'task_8_emi_threshold.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 9
    # --------------------------------------------------
    add_heading('Task 9: Loan Application Insights', 1, PRIMARY)
    add_body(
        'Approval rate is high but selective. Rejections are driven roughly equally by '
        'low credit score, insufficient income, and incomplete documentation — '
        'indicating that both credit quality and operational improvements can lift approval rates. '
        'Processing fees differ between approved and rejected applications.'
    )
    for fig in ['task_9_application_status.png', 'task_9_rejection_reason.png', 'task_9_processing_fee.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 10
    # --------------------------------------------------
    add_heading('Task 10: Recovery Effectiveness', 1, PRIMARY)
    add_body(
        'Overall recovery rate is critically low. Legal action provides marginal improvement. '
        'A tiered collections model and assignment of aged NPAs to Asset Reconstruction '
        'Companies (ARCs) is recommended to meaningfully improve outcomes.'
    )
    for fig in ['task_10_recovery_rate.png', 'task_10_recovery_legal.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 11
    # --------------------------------------------------
    add_heading('Task 11: Loan Disbursement Efficiency', 1, PRIMARY)
    add_body(
        'Processing time varies substantially across regions and loan purposes. '
        'Certain loan types create bottlenecks in the approval pipeline. '
        'Digital document submission and automated credit bureau checks could '
        'significantly reduce average processing days.'
    )
    for fig in ['task_11_disbursement_region.png', 'task_11_disbursement_purpose.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 12
    # --------------------------------------------------
    add_heading('Task 12: Revenue & Profitability Analysis', 1, PRIMARY)
    add_body(
        'Interest income varies across regions and loan purposes. '
        'Vehicle and business loan segments are the most profitable by purpose. '
        'Certain regions generate disproportionately higher revenue, '
        'creating geographic concentration risk in the revenue base.'
    )
    for fig in ['task_12_profit_purpose.png', 'task_12_profit_region.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 13
    # --------------------------------------------------
    add_heading('Task 13: Regional (Geospatial) Analysis', 1, PRIMARY)
    add_body(
        'Active loan distribution across regions is broadly balanced, reflecting '
        'healthy geographic diversification. Default rates differ meaningfully by region, '
        'suggesting that localised risk factors (economic conditions, collections capacity) '
        'drive regional default variance.'
    )
    for fig in ['task_13_geo_distribution.png', 'task_13_geo_default.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 14
    # --------------------------------------------------
    add_heading('Task 14: Default Trend Analysis', 1, PRIMARY)
    add_body(
        'Default counts fluctuate over the analysis period with no strong secular trend. '
        'Certain loan purposes (personal and education loans) exhibit higher average '
        'default amounts. Low-income borrowers show the highest default rate.'
    )
    for fig in ['task_14_default_trend.png', 'task_14_default_purpose.png', 'task_14_default_income.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 15
    # --------------------------------------------------
    add_heading('Task 15: Branch Efficiency — Feasibility Note', 1, PRIMARY)
    add_body(
        'Branch-level efficiency metrics such as per-branch processing time, rejection rate, '
        'and satisfaction score cannot be computed because the loans and applications datasets '
        'do not contain a BRANCH_ID column. This linkage gap is noted as a data improvement '
        'recommendation. Available branch metrics from branches.csv have been used where possible.'
    )
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 16
    # --------------------------------------------------
    add_heading('Task 16: Temporal Analysis', 1, PRIMARY)
    add_body(
        'Loan disbursement volumes are broadly steady month-on-month. '
        'Seasonal application patterns show moderate peaks in certain months. '
        'Regional default rates fluctuate over time, with some regions showing '
        'pronounced cycles that may align with economic seasonality.'
    )
    for fig in ['task_16_time_disbursement.png', 'task_16_time_seasonal_app.png',
                'task_16_time_seasonal_disb.png', 'task_16_time_default_region.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 17
    # --------------------------------------------------
    add_heading('Task 17: Repayment Behaviour Analysis', 1, PRIMARY)
    add_body(
        'The majority of customers make on-time repayments. A meaningful minority '
        'exhibit frequent or occasional late payment behaviour, which is a leading '
        'indicator of future default. Approval rates are broadly consistent across '
        'gender groups. Senior borrowers show slightly higher rejection rates.'
    )
    for fig in ['task_17_customer_behavior.png', 'task_17_customer_age.png', 'task_17_customer_gender.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 18
    # --------------------------------------------------
    add_heading('Task 18: Credit Risk Profiling', 1, PRIMARY)
    add_body(
        'Risk varies significantly across loan purposes — certain product categories '
        'carry materially higher default concentrations. Credit score tier is the '
        'strongest individual predictor of default risk, with low-score borrowers '
        'defaulting at substantially higher rates than high-score borrowers.'
    )
    for fig in ['task_18_risk_purpose.png', 'task_18_risk_credit.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 19
    # --------------------------------------------------
    add_heading('Task 19: Default Velocity Analysis', 1, PRIMARY)
    add_body(
        'Some loan categories default significantly faster after disbursement. '
        'Faster-defaulting segments represent higher risk exposure as recovery '
        'is more difficult for recently originated loans. '
        'Low credit score borrowers default earlier on average than high score borrowers.'
    )
    for fig in ['task_19_time_to_default_purpose.png', 'task_19_time_to_default_credit.png']:
        insert_figure(fig)
    doc.add_page_break()

    # --------------------------------------------------
    # TASK 20
    # --------------------------------------------------
    add_heading('Task 20: Transaction Behaviour Analysis', 1, PRIMARY)
    add_body(
        'Penalty transactions make up a large proportion of total transaction volume, '
        'reflecting systemic repayment stress across the portfolio. '
        'EMI transactions are the most common non-penalty type. '
        'Transaction behaviour patterns can serve as early warning signals for default risk.'
    )
    insert_figure('task_20_txn_type_dist.png')
    doc.add_page_break()

    # --------------------------------------------------
    # STRATEGIC RECOMMENDATIONS
    # --------------------------------------------------
    add_heading('Strategic Recommendations', 1, PRIMARY)

    recommendations = [
        ('Reduce Loan Defaults',
         'Enforce a minimum credit score threshold for unsecured products. '
         'Implement an early warning system at 3, 6, and 12 months post-disbursement. '
         'Cap EMI-to-income ratio at 40% at origination.'),
        ('Improve Recovery Rates',
         'Deploy a tiered collections model: automated reminders → tele-collections → '
         'field visits → legal/ARC referral. Assign NPAs older than 360 days to ARCs.'),
        ('Accelerate Processing',
         'Introduce digital document submission and automated bureau API checks '
         'to target a 30-day average processing cycle.'),
        ('Tackle the Overdue Portfolio',
         'Triage the overdue book immediately. Offer EMI restructuring and '
         'temporary moratorium programs to prevent overdue loans cascading to default.'),
        ('Grow Profitable Segments',
         'Focus origination capacity on vehicle and business loans which yield '
         'highest interest income. Convert document-rejection cases via guided digital onboarding.'),
        ('Data Infrastructure',
         'Add BRANCH_ID to loans and applications datasets to enable branch-level '
         'analytics. Build an ML-based credit scorecard for improved underwriting precision.'),
    ]

    for rec_title, rec_detail in recommendations:
        p = doc.add_paragraph()
        bold_run = p.add_run(f'• {rec_title}: ')
        bold_run.bold = True
        bold_run.font.color.rgb = PRIMARY
        p.add_run(rec_detail)

    # --------------------------------------------------
    # SAVE
    # --------------------------------------------------
    out_path = os.path.join(REPORT_PATH, 'mani_hero_fincorp_report.docx')
    doc.save(out_path)
    print(f"\n  Report saved: {out_path}")
    return out_path
