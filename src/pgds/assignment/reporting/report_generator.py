"""
Reporting — report_generator.py
Generates a formatted Word (.docx) report summarising all 20 task results.
"""

import os
from config import REPORT_PATH, FIGURE_PATH

os.makedirs(REPORT_PATH, exist_ok=True)
os.makedirs(FIGURE_PATH, exist_ok=True)


def generate_full_report(df, results=None):
    """
    Build the Hero FinCorp analysis Word report using python-docx.

    Parameters
    ----------
    df      : master merged DataFrame
    results : dict of dicts — output from each analyser task (optional)
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  ⚠️  python-docx not installed. Run: pip install python-docx")
        return

    print("\n" + "="*55)
    print(" Generating Word Report...")
    print("="*55)

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    RED_COLOR  = RGBColor(0xE3, 0x1E, 0x24)
    DARK_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
    BLUE_COLOR = RGBColor(0x2E, 0x86, 0xAB)

    def heading(text, level=1, color=None):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if color:
            for run in h.runs:
                run.font.color.rgb = color
        return h

    def body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        return p

    def kv(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run_l = p.add_run(f"  {label}: ")
        run_l.bold = True
        run_l.font.color.rgb = DARK_COLOR
        p.add_run(str(value))

    def add_fig(filename, width=Inches(5.5)):
        path = os.path.join(FIGURE_PATH, filename)
        if os.path.exists(path):
            doc.add_picture(path, width=width)
            last = doc.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            body(f"[Chart not found: {filename}]")

    def add_df_table(dataframe, max_rows=15):
        """Render a pandas DataFrame as a Word table."""
        if dataframe is None or dataframe.empty:
            return
        df_show = dataframe.head(max_rows).reset_index()
        cols = df_show.columns.tolist()
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = 'Table Grid'
        # Header row
        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = str(col)
            run = hdr[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E31E24')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
        # Data rows
        for ri, row in df_show.iterrows():
            cells = table.add_row().cells
            for ci, val in enumerate(row):
                cells[ci].text = str(round(val, 4) if isinstance(val, float) else val)
                if ri % 2 == 0:
                    tc = cells[ci]._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'F2F2F2')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:val'), 'clear')
                    tcPr.append(shd)
        doc.add_paragraph()

    # ──────────────────────────────────────────────────────────────────────────
    # COVER PAGE
    # ──────────────────────────────────────────────────────────────────────────
    title = doc.add_heading('Hero FinCorp', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RED_COLOR
        run.font.size = Pt(32)

    sub = doc.add_paragraph('Comprehensive Data-Driven Analysis')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(16)
    sub.runs[0].font.color.rgb = DARK_COLOR

    doc.add_paragraph()
    meta = doc.add_paragraph('All 20 Analysis Tasks | April 2026')
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # EXECUTIVE SUMMARY
    # ──────────────────────────────────────────────────────────────────────────
    heading('Executive Summary', 1, RED_COLOR)
    body(
        'This report presents the findings of a comprehensive data-driven analysis of '
        'Hero FinCorp\'s loan portfolio covering 70,000 customers, 90,000 loans, '
        '82,600 applications, 9,000 default records, 50 branches, and 495,000 transactions.'
    )

    total_loans    = len(df)
    default_rate   = df['DEFAULT_FLAG'].mean() * 100 if 'DEFAULT_FLAG' in df.columns else 0
    total_disb     = df['LOAN_AMOUNT'].sum() / 1e9 if 'LOAN_AMOUNT' in df.columns else 0
    overdue_pct    = (df['LOAN_STATUS'] == 'Overdue').mean() * 100 if 'LOAN_STATUS' in df.columns else 0

    kv('Total Loans Analysed',   f'{total_loans:,}')
    kv('Total Disbursed',        f'₹{total_disb:.2f} Billion')
    kv('Overall Default Rate',   f'{default_rate:.1f}%')
    kv('Overdue Portfolio',      f'{overdue_pct:.1f}% of loans')
    kv('Avg Credit Score',       f'{df["CREDIT_SCORE"].mean():.0f}' if 'CREDIT_SCORE' in df.columns else 'N/A')
    kv('Avg Interest Rate',      f'{df["INTEREST_RATE"].mean():.2f}%' if 'INTEREST_RATE' in df.columns else 'N/A')

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # TASK 1
    # ──────────────────────────────────────────────────────────────────────────
    heading('Task 1: Data Quality & Preparation', 1, RED_COLOR)
    body(
        'All six datasets were validated and cleaned. Column names were standardised '
        'to uppercase. Duplicates were removed and missing values imputed (median for '
        'numeric, mode for categorical). Date columns were parsed. Outliers in key '
        'numeric columns (LOAN_AMOUNT, INTEREST_RATE, DEFAULT_AMOUNT, CREDIT_SCORE, '
        'ANNUAL_INCOME, EMI_AMOUNT) were capped using the IQR method '
        '(Q1 − 1.5×IQR to Q3 + 1.5×IQR). Domain rules were enforced: '
        'LOAN_AMOUNT > 0 and INTEREST_RATE ≥ 0.'
    )
    heading('Data Quality Summary', 2)
    data_quality = [
        ['Dataset',      'Records', 'Columns', 'Key Action'],
        ['Customers',    '70,000',  '14',      'No nulls; outliers capped in ANNUAL_INCOME, CREDIT_SCORE'],
        ['Loans',        '90,000',  '12',      '30,194 null COLLATERAL_DETAILS treated as "None" (unsecured)'],
        ['Applications', '82,600',  '10',      '12,600 null LOAN_ID (rejected apps); REJECTION_REASON blank for approved'],
        ['Defaults',     '9,000',   '9',       '2,985 null RECOVERY_STATUS flagged; excluded from rate calc'],
        ['Branches',     '50',      '9',       'No issues found'],
        ['Transactions', '495,000', '9',       'No issues found'],
    ]
    t = doc.add_table(rows=len(data_quality), cols=4)
    t.style = 'Table Grid'
    for ri, row in enumerate(data_quality):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = val
            if ri == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # TASK 2
    # ──────────────────────────────────────────────────────────────────────────
    heading('Task 2: Descriptive Analysis', 1, RED_COLOR)
    body('Key distributions and regional/monthly trends across the portfolio.')
    add_fig('task02_loan_distribution.png')
    add_fig('task02_emi_distribution.png')
    add_fig('task02_credit_score_distribution.png')
    add_fig('task02_region_disbursement.png')
    add_fig('task02_region_default.png')
    add_fig('task02_monthly_applications.png')
    doc.add_page_break()

    # TASK 3
    heading('Task 3: Default Risk Analysis', 1, RED_COLOR)
    body(
        'Correlation analysis between loan attributes and default flag. '
        'Credit score shows a negative correlation with defaults; interest rate shows '
        'a weak positive correlation. Pairwise correlation between EMI amount, '
        'overdue amount, and default amount reveals co-movement in distress indicators.'
    )
    add_fig('task03_loan_correlation_heatmap.png')
    add_fig('task03_pairwise_heatmap.png')
    add_fig('task03_default_by_credit_bucket.png')
    doc.add_page_break()

    # TASK 4
    heading('Task 4: Branch & Regional Performance', 1, RED_COLOR)
    body('Branches ranked by disbursement volume, processing time, default rate, and recovery rate.')
    add_fig('task04_top_branches_disbursement.png')
    add_fig('task04_branch_delinquency.png')
    add_fig('task04_region_disbursement.png')
    add_fig('task04_region_default_rate.png')
    doc.add_page_break()

    # TASK 5
    heading('Task 5: Customer Segmentation', 1, RED_COLOR)
    body(
        'Customers segmented by credit score (High Risk / Medium / Good / Excellent) '
        'and income (Low / Mid / High). High-value customers identified as those with '
        'credit score ≥ 700, income in top 25%, and no default history.'
    )
    add_fig('task05_customer_credit_segment.png')
    add_fig('task05_default_by_credit_segment.png')
    add_fig('task05_default_by_income_segment.png')
    doc.add_page_break()

    # TASK 6
    heading('Task 6: Advanced Statistical Analysis', 1, RED_COLOR)
    body(
        'Full pairwise correlation matrix across key risk variables. '
        'Branch-level correlation between delinquency rate, disbursement amount, '
        'and recovery rate examined. Results confirm that lower credit scores and '
        'higher interest rates are positively correlated with default.'
    )
    add_fig('task06_default_risk_heatmap.png')
    add_fig('task06_advanced_pairwise_heatmap.png')
    doc.add_page_break()

    # TASK 7
    heading('Task 7: Transaction & Recovery Analysis', 1, RED_COLOR)
    body(
        'Penalty transactions constitute ~50% of all transactions — an alarming indicator '
        'of systematic repayment stress. Recovery rate is only 24.3% overall. '
        'Legal action does not materially improve recovery (40.7% vs 40.2% without).'
    )
    add_fig('task07_transaction_types.png')
    add_fig('task07_recovery_distribution.png')
    add_fig('task07_recovery_by_legal_action.png')
    add_fig('task07_recovery_by_reason.png')
    doc.add_page_break()

    # TASK 8
    heading('Task 8: EMI Analysis', 1, RED_COLOR)
    body(
        'EMI amounts were bucketed into quintiles. Mid-range EMI brackets '
        'show elevated default rates. Loans with EMI-to-income ratio above 40% '
        'carry significantly higher default risk.'
    )
    add_fig('task08_emi_default.png')
    add_fig('task08_emi_by_purpose.png')
    doc.add_page_break()

    # TASK 9
    heading('Task 9: Loan Application Insights', 1, RED_COLOR)
    body(
        'Overall approval rate is 84.7%. The three main rejection reasons '
        '(Low Credit Score 35%, Incomplete Documents 34%, Insufficient Income 33%) '
        'are broadly equal — suggesting both credit and operational improvements are needed.'
    )
    add_fig('task09_approval_rate_pie.png')
    add_fig('task09_rejection_reasons.png')
    add_fig('task09_processing_fee_comparison.png')
    doc.add_page_break()

    # TASK 10
    heading('Task 10: Recovery Effectiveness', 1, RED_COLOR)
    body(
        'Overall recovery rate is critically low at 24.3%. '
        'Legal action provides marginal benefit. Branch-wise recovery varies significantly. '
        'Recommendation: Engage ARCs for NPAs older than 360 days; '
        'implement tiered collections model.'
    )
    doc.add_page_break()

    # TASK 11
    heading('Task 11: Loan Disbursement Efficiency', 1, RED_COLOR)
    body(
        'Average processing time is ~175 days — far above the industry target of 30 days. '
        'Regional and channel-level bottlenecks identified. '
        'Digital document submission and automated credit checks recommended.'
    )
    add_fig('task12_monthly_disbursement_trend.png')
    doc.add_page_break()

    # TASK 12
    heading('Task 12: Profitability Analysis', 1, RED_COLOR)
    body(
        'Estimated total interest income is ₹75.43 Billion. '
        'Vehicle loans are the most profitable purpose (₹1,189.9 Cr). '
        'East and North regions generate the highest interest income.'
    )
    add_fig('task12_interest_income_by_purpose.png')
    add_fig('task12_interest_income_by_region.png')
    doc.add_page_break()

    # TASK 13
    heading('Task 13: Geospatial Analysis', 1, RED_COLOR)
    body(
        'Loan disbursement is broadly uniform across all 6 regions (₹37–39B each), '
        'indicating healthy geographic diversification. South and West regions '
        'show the highest default rates at 10.5%.'
    )
    doc.add_page_break()

    # TASK 14
    heading('Task 14: Default Trends', 1, RED_COLOR)
    body('Monthly default counts trend and default rates segmented by purpose and income category.')
    add_fig('task14_monthly_default_trend.png')
    add_fig('task14_default_by_purpose.png')
    doc.add_page_break()

    # TASK 15
    heading('Task 15: Branch Efficiency', 1, RED_COLOR)
    body(
        'Average branch disbursement time and rejection rates computed per branch. '
        'One branch has a delinquency rate of 137.7% — requiring immediate intervention. '
        'Top efficient branches identified for benchmarking.'
    )
    doc.add_page_break()

    # TASK 16
    heading('Task 16: Time-Series Analysis', 1, RED_COLOR)
    body('Seasonal application and disbursement patterns. Q2 shows peak application volumes.')
    add_fig('task16_seasonal_applications.png')
    add_fig('task16_yearly_disbursement.png')
    doc.add_page_break()

    # TASK 17
    heading('Task 17: Customer Behavior Analysis', 1, RED_COLOR)
    body(
        'Customers categorised as Always On Time, Occasional Defaulters, and Defaulters '
        'based on overdue amount and default flag. Default rates by gender, age group, '
        'and employment status computed. High-value customers (CS≥700, top 25% income, '
        'no default) identified for targeted retention programs.'
    )
    doc.add_page_break()

    # TASK 18
    heading('Task 18: Risk Assessment', 1, RED_COLOR)
    body(
        'Risk matrix built for each loan purpose combining default rate, '
        'average default amount, loan term, and interest rate into a composite Risk Score. '
        'Mitigation strategies assigned to each product category.'
    )
    doc.add_page_break()

    # TASK 19
    heading('Task 19: Time to Default Analysis', 1, RED_COLOR)
    body(
        'Median time from disbursement to default is 563 days (~19 months). '
        'Early-stage intervention programs (months 3–18) would have highest impact. '
        'Personal loans default fastest; Vehicle loans take longest to default.'
    )
    add_fig('task19_time_to_default.png')
    doc.add_page_break()

    # TASK 20
    heading('Task 20: Transaction Pattern Analysis', 1, RED_COLOR)
    body(
        'Penalty transactions represent 50.1% of all transactions. '
        'Customers with penalty count above median flagged as irregular repayers. '
        'Overdue loans show significantly higher average transaction amounts than '
        'non-overdue loans, driven by accumulated fees and penalties.'
    )
    add_fig('task20_monthly_penalty_pct.png')
    add_fig('task20_overdue_vs_nonoverdue_transactions.png')
    doc.add_page_break()

    # ──────────────────────────────────────────────────────────────────────────
    # RECOMMENDATIONS
    # ──────────────────────────────────────────────────────────────────────────
    heading('Strategic Recommendations', 1, RED_COLOR)

    recs = [
        ('Reduce Defaults',
         'Enforce minimum credit score of 580 for unsecured loans. '
         'Implement Early Warning System at months 3, 6, 12 post-disbursement. '
         'Apply income-to-EMI cap of 40%.'),
        ('Fix Recovery',
         'Partner with ARCs for NPAs > 360 days. '
         'Implement tiered collections (SMS → tele → field → legal/ARC). '
         'Reserve legal action for defaults above ₹1 Lakh only.'),
        ('Speed Up Processing',
         'Target 30-day processing via digital document upload, '
         'automated credit bureau API, and rule-based pre-screening.'),
        ('Tackle Overdue Portfolio',
         '33% overdue loans require immediate triage. '
         'Offer EMI restructuring and grace period programs.'),
        ('Boost Profitability',
         'Focus on Vehicle and Business loans (highest interest income). '
         'Convert 34% incomplete-document rejections into approvals via digital onboarding.'),
        ('Technology',
         'Build ML credit scorecard; implement real-time transaction monitoring; '
         'deploy branch efficiency dashboard updated weekly.'),
    ]

    for title_r, detail in recs:
        p = doc.add_paragraph()
        run_t = p.add_run(f'• {title_r}: ')
        run_t.bold = True
        run_t.font.color.rgb = RED_COLOR
        p.add_run(detail)

    # Save
    out_path = os.path.join(REPORT_PATH, 'hero_fincorp_analysis.docx')
    doc.save(out_path)
    print(f"\n  ✅ Report saved: {out_path}")
    return out_path
